"""
Document indexer for the Job Hunting U RAG chatbot.

Reads all .docx files in the knowledge base directory, splits them into
overlapping word-chunks tagged with source + section metadata, embeds them
locally with sentence-transformers, and stores them in a persistent
ChromaDB collection.

Run this once to build the vector DB. Re-run any time the source
documents change — it fully rebuilds the collection.
"""

import os
import sys
from pathlib import Path

import chromadb
import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer

load_dotenv()

KNOWLEDGE_BASE_DIR = Path(
    os.getenv("KNOWLEDGE_BASE_DIR", "../../knowledge_base_jhu")
).resolve()
CHROMA_DB_PATH = os.getenv("CHROMA_DB_PATH", "./jhu_vectordb")
CHROMA_COLLECTION_NAME = os.getenv("CHROMA_COLLECTION_NAME", "jhu_knowledge")
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"

CHUNK_SIZE_WORDS = 400
CHUNK_OVERLAP_WORDS = 50
CHUNK_STEP_WORDS = CHUNK_SIZE_WORDS - CHUNK_OVERLAP_WORDS


def iter_block_items(document):
    """Yield each Paragraph and Table in the document in true reading
    order (python-docx exposes .paragraphs and .tables as separate,
    unordered collections by default, which silently loses tables that
    appear alongside headings/paragraphs — e.g. a pricing table)."""
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def table_to_text_lines(table: Table) -> list[str]:
    """Flatten a table into readable 'Header: value' sentences, one per
    data row, using the first row as column headers."""
    rows = table.rows
    if not rows:
        return []

    headers = [cell.text.strip() for cell in rows[0].cells]
    lines = []
    for row in rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        parts = [f"{h}: {v}" for h, v in zip(headers, values) if h and v]
        if parts:
            lines.append(". ".join(parts))
    return lines


def extract_tagged_words(file_path: Path) -> list[tuple[str, str]]:
    """Read a .docx and return a flat list of (word, section) tuples.

    'section' tracks the most recent Heading 1 / Heading 2 text seen,
    so every word can be traced back to the part of the document it
    came from even after chunking. Tables are flattened into sentences
    and tagged with whatever section they appear under, same as text.
    """
    document = docx.Document(str(file_path))
    tagged_words: list[tuple[str, str]] = []
    current_heading1 = ""
    current_heading2 = ""

    for block in iter_block_items(document):
        if isinstance(block, Table):
            section = current_heading2 or current_heading1 or "General"
            for line in table_to_text_lines(block):
                for word in line.split():
                    tagged_words.append((word, section))
            continue

        paragraph = block
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else "Normal"

        if style_name == "Heading 1":
            current_heading1 = text
            current_heading2 = ""
            continue
        if style_name == "Heading 2":
            current_heading2 = text
            continue

        section = current_heading2 or current_heading1 or "General"
        for word in text.split():
            tagged_words.append((word, section))

    return tagged_words


def chunk_tagged_words(
    tagged_words: list[tuple[str, str]], source_name: str
) -> list[dict]:
    """Slide a 400-word window (50-word overlap) over the document,
    producing chunk dicts with text + metadata."""
    chunks = []
    total_words = len(tagged_words)
    if total_words == 0:
        return chunks

    start = 0
    chunk_index = 0
    while start < total_words:
        end = min(start + CHUNK_SIZE_WORDS, total_words)
        window = tagged_words[start:end]
        words = [w for w, _ in window]
        # Use the section of the first word in the window as the chunk's
        # primary label — good enough for citation purposes.
        section = window[0][1]

        chunks.append(
            {
                "text": " ".join(words),
                "source": source_name,
                "section": section,
                "chunk_index": chunk_index,
            }
        )

        chunk_index += 1
        if end == total_words:
            break
        start += CHUNK_STEP_WORDS

    return chunks


def main() -> None:
    if not KNOWLEDGE_BASE_DIR.exists():
        print(f"ERROR: knowledge base directory not found: {KNOWLEDGE_BASE_DIR}")
        sys.exit(1)

    docx_files = sorted(
        f
        for f in KNOWLEDGE_BASE_DIR.glob("*.docx")
        if not f.name.startswith("~$")  # skip Word lock files
    )

    if not docx_files:
        print(f"ERROR: no .docx files found in {KNOWLEDGE_BASE_DIR}")
        sys.exit(1)

    print(f"Found {len(docx_files)} document(s) in {KNOWLEDGE_BASE_DIR}\n")

    all_chunks: list[dict] = []
    for file_path in docx_files:
        print(f"Processing: {file_path.name} ...")
        try:
            tagged_words = extract_tagged_words(file_path)
            file_chunks = chunk_tagged_words(tagged_words, file_path.name)
            all_chunks.extend(file_chunks)
            print(f"  -> {len(file_chunks)} chunks created")
        except Exception as exc:
            print(f"  !! failed to process {file_path.name}: {exc}")

    if not all_chunks:
        print("ERROR: no chunks were produced. Aborting.")
        sys.exit(1)

    print(f"\nTotal chunks across all documents: {len(all_chunks)}")

    print(f"\nLoading embedding model '{EMBEDDING_MODEL_NAME}' ...")
    embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)

    print("Embedding chunks (this may take a moment) ...")
    texts = [c["text"] for c in all_chunks]
    embeddings = embedding_model.encode(
        texts, show_progress_bar=True, convert_to_numpy=True
    ).tolist()

    print(f"\nConnecting to ChromaDB at {CHROMA_DB_PATH} ...")
    client = chromadb.PersistentClient(path=CHROMA_DB_PATH)

    # Rebuild the collection from scratch each run so stale chunks from
    # removed/edited documents never linger.
    try:
        client.delete_collection(CHROMA_COLLECTION_NAME)
        print(f"Deleted existing collection '{CHROMA_COLLECTION_NAME}'")
    except Exception:
        pass

    collection = client.create_collection(name=CHROMA_COLLECTION_NAME)

    ids = [f"{c['source']}::{c['chunk_index']}" for c in all_chunks]
    metadatas = [
        {
            "source": c["source"],
            "section": c["section"],
            "chunk_index": c["chunk_index"],
        }
        for c in all_chunks
    ]

    collection.add(
        ids=ids,
        documents=texts,
        embeddings=embeddings,
        metadatas=metadatas,
    )

    print(
        f"\nDone. Stored {len(all_chunks)} chunks in collection "
        f"'{CHROMA_COLLECTION_NAME}' at {CHROMA_DB_PATH}"
    )


if __name__ == "__main__":
    main()
